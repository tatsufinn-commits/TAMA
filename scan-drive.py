#!/usr/bin/env python3
"""
TAMA Google Drive Remote In-Memory Scanner (Zero-Disk-Footprint)
--------------------------------------------------------------
Allows AI agents and students to search, inspect, and extract text from
the TAMA Google Drive Resource Hub without saving binary files to disk.

Storage Impact: 0 BYTES on persistent disk (Streams directly into RAM).
"""

import sys
import io
import urllib.request
import urllib.parse
import json

# Curated catalog of all files in the TAMA Google Drive Resource Hub
DRIVE_CATALOG = [
    # 1. HOA 1-4 Series & Lecture Slides
    {"id": "11lVCrLJqDi_ILApRyk0GnRt-B03u9hw_", "name": "Deptl-Exam_ARIDBE2025-AR133-1P-1st-Tsem.pdf", "type": "pdf", "category": "Exam", "week": "Deptal Exam AY 2025-2026"},
    {"id": "1a3Zq8ZNbB0f9RtjwTZz2wfMuaIHbEM8m", "name": "HOAX-1st-Take.pdf", "type": "pdf", "category": "HOA", "week": "Comprehensive HOA Reviewer"},
    {"id": "1NW5JipGcHiVd9c6qWuG0Nk_BsAgq2wdN", "name": "W1_1 - Pre-Historic Architecture & Part 1 Egyptian Architecture.docx", "type": "docx", "category": "HOA1", "week": "Week 1"},
    {"id": "19mViDZ8rfbTVShXd1iDeknotGeHcqGA8", "name": "W1_1 - Pre-Historic Architecture & Part 1 Egyptian Architecture.pdf", "type": "pdf", "category": "HOA1", "week": "Week 1"},
    {"id": "1FYkpE1U87WKDXEkSA5VtxyYch2iYyFG3", "name": "W1_2 Egyptian Architecture.pdf", "type": "pdf", "category": "HOA1", "week": "Week 1"},
    {"id": "1mQW3UH8hRTUjxCowbNvGfdVxBH18m4dH", "name": "W1_3 - Asia Babylonian & Asiatic Sumerian Architecture.pdf", "type": "pdf", "category": "HOA1", "week": "Week 1"},
    {"id": "1rAmLMo1Nibb0lAXbXYR77KoB6o3HlxyZ", "name": "W1_4 - West Asiatic Assyrian Architecture.pdf", "type": "pdf", "category": "HOA1", "week": "Week 1"},
    {"id": "1TkYdnF9jkW2SJI3IcwzGQrldj1CgP5YJ", "name": "W2_5 - Minoan Architectural Influences and Character.pdf", "type": "pdf", "category": "HOA1", "week": "Week 2"},
    {"id": "1wLLRXmyDmisG7L1x-tsC7OBdlWv3PFSV", "name": "W2_6 - History of Mycenaean Architecture.pdf", "type": "pdf", "category": "HOA1", "week": "Week 2"},
    {"id": "1qi6EnVzLRCuzzmdBQF2YjgOnBGDWpEza", "name": "W2_7 - Intro to Classical Greek Archi, Other Greek Structures, and Greek Orders.pdf", "type": "pdf", "category": "HOA1", "week": "Week 2"},
    {"id": "1Js11HMmJEkIsE7e43t1F4Ft1Xah8HJ4p", "name": "W2_8 - Part 1 Roman Architecture Historical Development.pdf", "type": "pdf", "category": "HOA1", "week": "Week 2"},
    {"id": "1jMP6B7hCDKCkGiIKlRr7yN5-IiKtMS5z", "name": "W2_9 - Part 2 Etruscan and Early Roman Period.pdf", "type": "pdf", "category": "HOA1", "week": "Week 2"},
    {"id": "1SKYPOkjlV4R9gGjKOb46Mxt6NoVtNNpV", "name": "W2_10 - Part 3 Imperial Rome Architecture.pdf", "type": "pdf", "category": "HOA1", "week": "Week 2"},
    {"id": "1crgoHW_keoHa2-K-YIKDz5AwmOPABq39", "name": "W2_11 - History of Byzantine Empire.pdf", "type": "pdf", "category": "HOA2", "week": "Week 2"},
    {"id": "195Rnt-CWjAq3w_eExXoU2F2f2jNIOVOk", "name": "W2_12 - Architectural Influences of Byzantine Architecture.pdf", "type": "pdf", "category": "HOA2", "week": "Week 2"},
    {"id": "1nVxkUCIM6URyr8THCMvzMuLIkzIHqYLw", "name": "W2_13 - Byzantine Structures and Its Details.pdf", "type": "pdf", "category": "HOA2", "week": "Week 2"},
    {"id": "1g_Ku6z5NW3W9CFyf6PIAVRZidOHftovC", "name": "W2_14 - Brief History of Christianity Structures and Details of Early Christian Church.pdf", "type": "pdf", "category": "HOA2", "week": "Week 2"},
    {"id": "1x7voVuwzMk33-MfpAKkOTdv1cnknxHbz", "name": "W2_15 - Romanesque Architecture.pdf", "type": "pdf", "category": "HOA2", "week": "Week 2"},
    {"id": "1QRrEIzbYAxlVECHST5UK-kgAH8Oif7A_", "name": "W3_16 - Gothic Architecture & Gothic Vocabulary.pdf", "type": "pdf", "category": "HOA2", "week": "Week 3"},
    {"id": "1MYNjaifeTQBGQu80xlqyTuY3mpxSZwr8", "name": "W3_17 - Renaissance Overview & Renaissance in Italy.pdf", "type": "pdf", "category": "HOA3", "week": "Week 3"},
    {"id": "1-aBlK7PVkR-A4zhvzO6RqUUAhUX5Gels", "name": "W3_18 - RENAISSANCE IN FRANCE & SPANISH RENAISSANCE.pdf", "type": "pdf", "category": "HOA3", "week": "Week 3"},
    {"id": "1QM9tJNCsPjfsicWF6-uxKZGx6VsOTalw", "name": "W3_19 - British Renaissance & Russian Renaissance Architecture.pdf", "type": "pdf", "category": "HOA3", "week": "Week 3"},
    {"id": "1KHAyC1a-QoQKCrztTQxsaAhGXwvpm2QW", "name": "W3_20 - Gothic Revival & Greek Revival.pdf", "type": "pdf", "category": "HOA3", "week": "Week 3"},
    {"id": "1dw7A6RKpNuJR_iZr7pvSNTqykARaw80Z", "name": "W3_21 - Romanesque Revival & Neoclassical Style.pdf", "type": "pdf", "category": "HOA3", "week": "Week 3"},
    {"id": "1qQDRA8ThyJnTESyrlrsw2DK9vKrHI_QO", "name": "W3_22 - Art Deco & Art Nouveau.pdf", "type": "pdf", "category": "HOA3", "week": "Week 3"},
    {"id": "1crzeF0Ns4fjodv09R16v9LLvwODNNwOa", "name": "W3_23- Introduction to Modern Architecture and The Star Architects.pdf", "type": "pdf", "category": "HOA4", "week": "Week 3"},
    {"id": "1B5PqBFSY5Ea7K8KXZk6DXs413C5-oHaf", "name": "W4_24-India Architecture.pdf", "type": "pdf", "category": "HOA Asian", "week": "Week 4"},
    {"id": "1NlkXZftaNhM5jZfIyNjyTD5zAOSY77q0", "name": "W4_25 - Chinese Architecture.pdf", "type": "pdf", "category": "HOA Asian", "week": "Week 4"},
    {"id": "1ntUzsthibjjSW9wg-U2q234BEC9aKte0", "name": "W4_26 - Korean Architecture.pdf", "type": "pdf", "category": "HOA Asian", "week": "Week 4"},
    {"id": "1wlG6X9UWgqlxdRQIS8bw7Lyf7xyj7thc", "name": "W4_27 - Japanese Architecture.pdf", "type": "pdf", "category": "HOA Asian", "week": "Week 4"},
    {"id": "1OrdNDUH0KPOWi5MDHFvy3Cc26m4BBYgn", "name": "W4_28 - Thailand Architecture.pdf", "type": "pdf", "category": "HOA Asian", "week": "Week 4"},
    {"id": "1vueyPHToTIosF79IfZTO5QSKv1rxeTtL", "name": "W4_29 - Indonesian Architecture.pdf", "type": "pdf", "category": "HOA Asian", "week": "Week 4"},
    {"id": "17W5cQRH8EMu1o1T6K58P3WzwcF3PuhZU", "name": "W4_30 - Malaysian Architecture.pdf", "type": "pdf", "category": "HOA Asian", "week": "Week 4"},
    {"id": "1Oj-YnbL93XZzkFtgI6aTleZA2YnHIutF", "name": "W5_31 - Badjao & Samal Architecture.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 5"},
    {"id": "1KgJQM4ffaIEP97f2i7oHnY-fceJwZZXX", "name": "W5_32 - Tausug and Maranao Architecture.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 5"},
    {"id": "1sV94znWMYG05Ip2OVdyw_6nvambbTTyQ", "name": "W5_33-IFUGAO&BONTOC.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 5"},
    {"id": "1rMW-06um_aUIJ-U2slqHVj5LKqocLYXu", "name": "W5_34 - Isneg of Apayao and Kalinga.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 5"},
    {"id": "1mb2yP6TE_YVMqMm_tRrsHCKR-1ZjJ_TT", "name": "W5_35 - Ivatan Architecture & Panay Architecture.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 5"},
    {"id": "1P9jbV6Qz-0oyPl7PeK4c-V7WlpYrueYb", "name": "W6_36 - Bahay na Bato & Church Architecture during Spanish Colonial.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 6"},
    {"id": "1hCDAvoFQtHVBiXM88usGWDf548Fmk1T3", "name": "W6_37 - Settlements during the Spanish Colonial Period.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 6"},
    {"id": "13Xhi-OdmvDmDjxSUQoHuKW5J79Qh2mvQ", "name": "W6_38 - Spanish Colonial Fortresses & American Colonial Period.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 6"},
    {"id": "17mNxRMACI9zfEGDcSk0ErZJGDE5OWxXu", "name": "W6_39 - ARCHITECTURE IN THE PHILIPPINES MARTIAL LAW PERIOD.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 6"},
    {"id": "1m1szNVOD_vCOotklzay8L5o68oPKcT6F", "name": "W6_40 - EARLY MODERN ARCHITECTURE IN THE PHILIPPINES EARLY REPUBLIC.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 6"},
    {"id": "1wDK8D0mrb0IWQ7FwkHCF4D1Lf-s_pRtg", "name": "W6_41 - LAW ON HERITAGE CONSERVATION AND HERITAGE CONSERVATION.pdf", "type": "pdf", "category": "HOA PH", "week": "Week 6"},
    {"id": "1t4-gkkr5GJptodss_07l0jSoCnMuLSWt", "name": "W10_42-Iconic Architects Part 1.pdf", "type": "pdf", "category": "Monographs", "week": "Week 10"},
    {"id": "1YT85oYBMmaNTl2vimUPWVqtHgaUNtvzV", "name": "W11_43 - Iconic Architects Part 2.pdf", "type": "pdf", "category": "Monographs", "week": "Week 11"},
    {"id": "1rbM-fRCFawlvuirBmBytLdPjaS-IbJPh", "name": "W12_44-Pritzker Architecture Prize Part 1.pdf", "type": "pdf", "category": "Awards", "week": "Week 12"},
    {"id": "1fQ1qcyVqk1H2kbLsozK9BxumbRFfuW5b", "name": "W13_45-Pritzker Architecture Prize Part 2.pdf", "type": "pdf", "category": "Awards", "week": "Week 13"},
    # 2. TOA Reviewers
    {"id": "1CPJ6jG0sWDWKqERJCNII1cCkhrOBkuet", "name": "Principles-of-Design-and-Ordering-Principles.pdf", "type": "pdf", "category": "TOA", "week": "TOA Ching"},
    {"id": "1y5C2ALAoih5uBc-TCCHpBehdAZtXSg9p", "name": "PRITZKER AWARDEESS.pdf", "type": "pdf", "category": "TOA", "week": "Pritzker Profiles"},
    {"id": "1hfWDnKIfbiL-5-Uvnnxi9JuqdhJLBFnl", "name": "TOA-X REVIEWER_merged.pdf", "type": "pdf", "category": "TOA", "week": "TOA Mock Exam Bank"},
    # 3. Books
    {"id": "1BNzyPCPwizvzkp5kRsW3t7Sa7S-BGJYd", "name": "Planning-and-Designers-Handbook-2nd-Edition-by-Max-B-Fajardo-Jr.pdf", "type": "pdf", "category": "Textbook", "week": "Planning"},
    {"id": "1QDFogw3YLHbYBJlIp5hypFQoI1rgGu-U", "name": "Banister Fletcher - A History of Architecture on the Comparative Method.pdf", "type": "pdf", "category": "Textbook", "week": "HOA Classical"}
]

def list_catalog():
    print("\n📚 === TAMA GOOGLE DRIVE REMOTE CLOUD CATALOG ===")
    print(f"Total Available Files: {len(DRIVE_CATALOG)} (Zero Local Disk Footprint)\n")
    print(f"{'INDEX':<6} {'CATEGORY':<12} {'WEEK / DOMAIN':<20} {'FILE NAME'}")
    print("-" * 80)
    for i, item in enumerate(DRIVE_CATALOG, 1):
        print(f"[{i:02d}]   {item['category']:<12} {item['week']:<20} {item['name']}")
    print("-" * 80)
    print("\n👉 Usage: python3 scan-drive.py search <keyword> | python3 scan-drive.py read <file_id_or_index>\n")

def search_catalog(query):
    q = query.lower()
    results = [item for item in DRIVE_CATALOG if q in item['name'].lower() or q in item['category'].lower() or q in item['week'].lower()]
    print(f"\n🔍 Search results for '{query}' ({len(results)} found):")
    print("-" * 80)
    for item in results:
        print(f"• [{item['category']}] ({item['week']}) {item['name']}")
        print(f"  ID: {item['id']}")
    print("-" * 80)

def stream_and_scan(target, max_pages=5):
    """Streams a file into memory buffer (RAM) and extracts text without writing to disk."""
    # Resolve index or file ID
    file_info = None
    if target.isdigit() and 1 <= int(target) <= len(DRIVE_CATALOG):
        file_info = DRIVE_CATALOG[int(target) - 1]
    else:
        for item in DRIVE_CATALOG:
            if item['id'] == target or item['name'] == target:
                file_info = item
                break
    
    file_id = file_info['id'] if file_info else target
    file_name = file_info['name'] if file_info else f"Drive-File-{target}"

    print(f"\n🌐 Streaming '{file_name}' directly into memory (0 bytes on disk)...")
    url = f"https://drive.google.com/uc?id={file_id}&export=download"
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            content_bytes = resp.read()
            in_memory_buffer = io.BytesIO(content_bytes)
            size_kb = len(content_bytes) / 1024
            print(f"✅ Streamed {size_kb:.1f} KB into RAM buffer successfully.")
            
            # Parse text based on file type
            if file_name.endswith('.pdf'):
                try:
                    import pypdf
                    reader = pypdf.PdfReader(in_memory_buffer)
                    total_p = len(reader.pages)
                    print(f"📖 PDF Total Pages: {total_p} | Extracting first {min(max_pages, total_p)} pages...\n")
                    print("=" * 80)
                    for p_num in range(min(max_pages, total_p)):
                        text = reader.pages[p_num].extract_text()
                        print(f"\n--- [PAGE {p_num + 1} OF {total_p}] ---")
                        print(text.strip())
                    print("\n" + "=" * 80)
                    print("✅ In-memory scan complete. 0 bytes written to persistent disk.")
                except ImportError:
                    print("pypdf not available; install via pip install pypdf")
            
            elif file_name.endswith('.docx'):
                try:
                    import docx
                    doc = docx.Document(in_memory_buffer)
                    print(f"📖 Word Document: {len(doc.paragraphs)} paragraphs | Extracting text...\n")
                    print("=" * 80)
                    for i, p in enumerate(doc.paragraphs[:30]):
                        if p.text.strip():
                            print(f"{p.text.strip()}")
                    print("\n" + "=" * 80)
                except ImportError:
                    print("python-docx not available")
                    
    except Exception as e:
        print(f"❌ Error streaming file from Google Drive: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        list_catalog()
    elif sys.argv[1] == "search" and len(sys.argv) > 2:
        search_catalog(" ".join(sys.argv[2:]))
    elif sys.argv[1] in ["scan", "read", "inspect"] and len(sys.argv) > 2:
        stream_and_scan(sys.argv[2])
    else:
        list_catalog()
